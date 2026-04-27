#!/usr/bin/env python3
from docx import Document
import sys

# Load DOCX file
doc_path = r'C:\Users\OSCARPACK\Downloads\EMR\medsync_full_spec.docx'
doc = Document(doc_path)

# Extract all text
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text)

# Also extract from tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text)

output = '\n'.join(full_text)

# Save to file
with open(r'C:\Users\OSCARPACK\Downloads\EMR\spec_extracted.txt', 'w', encoding='utf-8') as f:
    f.write(output)

print(f"Extracted {len(full_text)} sections")
print(f"Total length: {len(output)} characters")
print("\n=== First 3000 characters ===\n")
print(output[:3000])
